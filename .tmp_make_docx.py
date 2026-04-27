from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = "/home/wjttdbx/python_ekf/项目说明_近距离相对导航与制导仿真.docx"

doc = Document()

styles = doc.styles
styles['Normal'].font.name = '宋体'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(12)

for sname, size in [('Title', 20), ('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
    if sname in styles:
        styles[sname].font.name = '黑体' if sname != 'Title' else '黑体'
        styles[sname]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        styles[sname].font.size = Pt(size)

section = doc.sections[0]
section.top_margin = Pt(72)
section.bottom_margin = Pt(72)
section.left_margin = Pt(90)
section.right_margin = Pt(90)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('近距离相对导航与制导仿真项目说明')
r.bold = True
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('基于非线性相对动力学、EKF状态估计与SDRE反馈控制')
r.font.name = '宋体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r.font.size = Pt(12)

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('日期：2026年4月21日')
r.font.name = '宋体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r.font.size = Pt(12)

doc.add_page_break()

doc.add_paragraph('目录', style='Heading 1')
ptoc = doc.add_paragraph()
run = ptoc.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar)
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
run._r.append(instrText)
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'separate')
run._r.append(fldChar)
ptoc.add_run('右键目录并选择“更新域”即可生成目录。')
run = ptoc.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar)

doc.add_page_break()

doc.add_heading('1 背景与研究意义', level=1)
doc.add_paragraph(
    '随着在轨服务、空间碎片清理、交会对接、临近探测等任务需求的持续增长，面向非合作目标的近距离相对导航与制导已成为航天自主控制领域的重要研究方向。相较于合作目标，非合作目标通常缺乏主动配合信息，例如协同信标、稳定通信链路或高精度先验轨道信息，因此任务系统需要在量测受限、噪声存在以及目标状态不可完全观测的条件下，实现对目标相对运动状态的准确估计与稳定控制。'
)
doc.add_paragraph(
    '本项目围绕“相对导航—状态估计—闭环制导”技术链路，构建了基于非线性椭圆轨道相对运动模型、扩展卡尔曼滤波方法以及状态相关黎卡提方程反馈控制方法的一体化仿真框架。该框架能够描述椭圆参考轨道条件下的三维相对运动过程，并对不同初始几何关系、不同目标机动条件以及不同轨道平面关系下的闭环性能进行统一分析。'
)
doc.add_paragraph(
    '该工作可为近距离自主接近、相对导航算法验证、制导律性能评估以及后续工程实现提供基础支撑，同时也可作为多场景对比研究的平台，用于系统分析初始条件变化对估计精度、收敛能力和控制代价的影响。'
)

doc.add_heading('2 理论模型与方法', level=1)

doc.add_heading('2.1 动力学模型', level=2)
doc.add_heading('2.1.1 坐标系与状态定义', level=3)
doc.add_paragraph(
    '本文采用LVLH坐标系描述追踪器与目标器的相对运动，其中x轴为径向方向，y轴为沿轨方向，z轴为轨道法向方向。追踪器与目标器的绝对状态分别表示为位置与速度构成的六维向量，相对状态定义为两者状态之差。该定义能够直接服务于后续相对导航建模与反馈控制设计。'
)
doc.add_heading('2.1.2 非线性椭圆轨道相对运动模型', level=3)
doc.add_paragraph(
    '在椭圆参考轨道条件下，系统状态除两航天器的绝对状态外，还包含参考轨道真近点角。模型考虑了参考轨道半径变化、角速度变化以及LVLH坐标系中的非惯性项，从而形成完整的非线性相对运动描述。与圆轨道线性模型相比，该模型能够更真实地反映大偏心率轨道下的相对动力学特性，更适用于复杂工况分析。'
)
doc.add_paragraph(
    '为便于在线反馈控制求解，系统进一步构造为状态相关线性形式，即在每个时刻根据当前状态生成状态相关系数矩阵，将非线性动力学等效表示为随状态变化的线性系统。这种处理方式兼顾了模型精度与控制设计便利性，是实现SDRE方法的基础。'
)

doc.add_heading('2.2 量测模型', level=2)
doc.add_paragraph(
    '系统采用距离—方位角—仰角量测模型。距离由目标相对位置向量模长给出，方位角由径向与沿轨分量构成的平面几何关系确定，仰角由法向分量与距离之间的关系确定。该量测模型符合典型近距离相对观测场景下的角距联合测量特点。'
)
doc.add_paragraph(
    '考虑到方位角与仰角属于角度量，在滤波更新过程中需要对角度残差进行归一化处理，以避免跨越正负π边界时出现突变误差。此外，为支撑扩展卡尔曼滤波器实现，需要在先验状态处对量测方程进行线性化，构造相应的量测雅可比矩阵。'
)

doc.add_heading('2.3 状态估计方法', level=2)
doc.add_paragraph(
    '为获得目标的实时相对状态估计，系统采用扩展卡尔曼滤波方法。滤波器以相对位置和相对速度为状态变量，以距离、方位角和仰角为量测输入。预测阶段利用状态相关动力学矩阵和控制输入对状态及协方差进行传播；更新阶段利用当前量测值与预测量测值之间的残差完成状态修正。'
)
doc.add_paragraph(
    '在闭环过程中，滤波器输出的相对状态估计结果直接提供给制导控制器使用，从而形成“动力学传播—量测更新—反馈控制”的连续工作链路。该方法能够在存在量测噪声和模型不确定性的情况下保持较好的相对状态恢复能力。'
)

doc.add_heading('2.4 制导与SDRE求解', level=2)
doc.add_paragraph(
    '控制器采用基于状态相关黎卡提方程的反馈设计思路。其基本思想是在每个控制时刻，将当前非线性系统表示为与状态相关的线性二次型调节问题，通过在线求解相应的代数黎卡提方程，获得与当前状态一致的反馈增益矩阵，并据此生成控制指令。'
)
doc.add_paragraph(
    '相较于固定增益线性方法，SDRE方法能够更好地适应非线性轨道相对运动中的时变性与状态依赖性，尤其适用于椭圆轨道和大范围相对运动工况。工程实现中，为提高在线求解稳定性，可采用矩阵尺度平衡策略改善黎卡提方程数值条件；同时可通过“间隔求解、步间复用”的方式降低计算负担。'
)
doc.add_paragraph(
    '在本项目中，制导目标是使追踪器相对目标的距离逐步减小，并在满足约束条件下实现稳定接近。控制效果通过相对距离收敛、估计误差变化及控制输入幅值等指标进行综合评估。'
)

doc.add_heading('3 仿真情形设计', level=1)
doc.add_paragraph(
    '为系统评估不同初始条件和目标机动条件对算法性能的影响，本文设置多组对比场景，覆盖共面与非共面、目标有无机动、同轨不同位置与不同相位等典型情况。各场景均在统一动力学框架和统一滤波控制流程下进行仿真。'
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '场景编号'
hdr[1].text = '场景描述'
hdr[2].text = '主要变量'
hdr[3].text = '分析目的'
scenarios = [
    ('A', '同轨共面、目标无机动', '基准初始距离；共面；目标不施加控制', '验证基线收敛能力与稳态精度'),
    ('B', '同轨共面、目标小幅机动', '在基准场景上增加弱机动', '评估轻度机动对估计和控制的影响'),
    ('C', '同轨共面、目标中强机动', '提高机动强度或机动频次', '分析扰动增强后的鲁棒性'),
    ('D', '同轨不同相位、目标位于前方', '沿轨初始偏差为主', '分析相位差对收敛时间和控制代价的影响'),
    ('E', '同轨不同相位、目标位于后方', '与D互补的沿轨初始关系', '比较前后位形差异引起的性能变化'),
    ('F', '非共面、小夹角、目标无机动', '存在明显法向偏差', '验证三维收敛能力及法向控制需求'),
    ('G', '非共面、小夹角、目标有机动', '非共面条件下叠加目标机动', '考察复杂耦合工况下的稳定性'),
    ('H', '远距离初始条件', '增大初始相对距离', '分析远距离条件下的收敛域与控制代价')
]
for row in scenarios:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text

doc.add_paragraph(
    '上述场景设置能够从多个维度反映初始条件对算法性能的影响。其中，共面与非共面差异主要体现三维耦合效应；有无机动差异主要体现控制鲁棒性；同轨不同相位或不同初始位置差异主要体现几何关系变化对收敛时间、控制幅值和量测可观性的影响。'
)

doc.add_heading('4 结果分析', level=1)
doc.add_paragraph(
    '建议对各组场景从以下几个方面进行统一分析：第一，相对距离随时间变化情况，用于衡量是否能够实现稳定接近以及达到预定捕获阈值；第二，位置与速度估计误差变化情况，用于评价滤波器在不同工况下的估计精度；第三，控制输入大小及累计控制代价，用于衡量制导过程的执行负担；第四，创新序列与协方差变化情况，用于判断滤波过程的一致性与稳定性。'
)
doc.add_paragraph(
    '从总体规律看，在同轨共面且目标无机动的基准场景下，系统通常表现出最快的距离收敛速度和较小的控制需求，估计误差也更易保持平稳。当目标施加机动后，量测残差和估计误差会在短时内放大，控制器需要输出更高幅值的修正指令，导致收敛过程延长。'
)
doc.add_paragraph(
    '在非共面场景中，法向相对运动引入额外耦合关系，使得三维控制与估计难度增加，特别是在初始法向偏差较大时更为明显。同轨不同相位场景则表明，沿轨方向初始偏差往往决定主要收敛时间尺度，并对控制能量消耗具有直接影响。对于远距离初始条件，系统虽然仍可保持总体收敛趋势，但模型非线性增强、量测几何变化和控制修正幅值增大，都会提高闭环设计的挑战。'
)
doc.add_paragraph(
    '因此，在工程应用中，应针对不同任务阶段采取差异化参数整定策略。远距离阶段更应关注稳定性与可观测性，近距离阶段则更强调高精度估计与平滑控制。同时，可根据目标机动强度和轨道几何关系，对滤波噪声参数、状态权重矩阵及控制惩罚矩阵进行针对性调整，以提升系统整体性能。'
)

doc.add_heading('5 结论', level=1)
doc.add_paragraph(
    '本文形成了一套面向近距离相对导航与制导问题的仿真说明框架，围绕非线性椭圆轨道相对运动建模、角距量测建模、扩展卡尔曼滤波状态估计以及SDRE反馈控制方法，对系统组成和技术路线进行了系统梳理。通过设置多组具有代表性的仿真场景，可较全面地分析目标机动、轨道平面关系和初始几何条件变化对闭环性能的影响。'
)
doc.add_paragraph(
    '后续可在本说明基础上进一步补充仿真图、定量统计表和参数敏感性分析结果，形成完整的项目技术报告。'
)

doc.save(OUT)
print(OUT)
